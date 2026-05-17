"""Edit existing DOCX files — modify, delete, search/replace, insert."""

from docx import Document
from docx.shared import Pt


def edit_docx_paragraph(file_path: str, index: int, new_text: str, output_path: str | None = None) -> str:
    """Edit a paragraph at the given index, preserving style."""
    doc = Document(file_path)
    if index < 0 or index >= len(doc.paragraphs):
        raise ValueError(f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})")

    para = doc.paragraphs[index]
    # Preserve style
    style = para.style
    # Clear and rewrite
    para.clear()
    para.add_run(new_text)
    para.style = style

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Paragraph {index} updated. Saved to {save_path}"


def delete_docx_paragraph(file_path: str, index: int, output_path: str | None = None) -> str:
    """Delete a paragraph at the given index."""
    doc = Document(file_path)
    if index < 0 or index >= len(doc.paragraphs):
        raise ValueError(f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})")

    para = doc.paragraphs[index]
    p_element = para._element
    p_element.getparent().remove(p_element)

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Paragraph {index} deleted. Saved to {save_path}"


def search_replace_docx(file_path: str, search_text: str, replace_text: str, output_path: str | None = None) -> str:
    """Search and replace text throughout the document."""
    doc = Document(file_path)
    count = 0

    for para in doc.paragraphs:
        if search_text in para.text:
            for run in para.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)
                    count += 1

    # Also search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if search_text in para.text:
                        for run in para.runs:
                            if search_text in run.text:
                                run.text = run.text.replace(search_text, replace_text)
                                count += 1

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Replaced {count} occurrence(s) of '{search_text}'. Saved to {save_path}"


def insert_docx_paragraph(file_path: str, index: int, text: str, style: str = "Normal", output_path: str | None = None) -> str:
    """Insert a new paragraph at the given index."""
    doc = Document(file_path)
    if index < 0 or index > len(doc.paragraphs):
        raise ValueError(f"Index {index} out of range (0-{len(doc.paragraphs)})")

    # Get the element to insert before
    if index < len(doc.paragraphs):
        ref_element = doc.paragraphs[index]._element
    else:
        ref_element = None

    from docx.oxml.ns import qn
    from copy import deepcopy

    new_para = doc.add_paragraph(text, style=style)
    new_element = new_para._element

    # Move to correct position
    body = doc.element.body
    body.remove(new_element)
    if ref_element is not None:
        ref_element.addprevious(new_element)
    else:
        body.append(new_element)

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Paragraph inserted at index {index}. Saved to {save_path}"

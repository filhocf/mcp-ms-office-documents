"""Tests for pdf_tools — PDF generation and conversion."""

import pytest
from unittest.mock import patch
from docx import Document

from pdf_tools.base_pdf_tool import markdown_to_pdf, docx_to_pdf, MarkdownPDF


class TestMarkdownPDF:
    def test_render_headings(self):
        pdf = MarkdownPDF()
        pdf.render_markdown("# Title\n## Subtitle\n### Section")
        buffer = pdf.output()
        assert len(buffer) > 0

    def test_render_paragraphs(self):
        pdf = MarkdownPDF()
        pdf.render_markdown("This is a paragraph.\n\nThis is another.")
        buffer = pdf.output()
        assert len(buffer) > 0

    def test_render_lists(self):
        pdf = MarkdownPDF()
        pdf.render_markdown("- Item 1\n- Item 2\n* Item 3\n1. Numbered")
        buffer = pdf.output()
        assert len(buffer) > 0

    def test_render_horizontal_rule(self):
        pdf = MarkdownPDF()
        pdf.render_markdown("Above\n---\nBelow")
        buffer = pdf.output()
        assert len(buffer) > 0

    def test_render_mixed_content(self):
        md = """# Report

## Introduction

This is the intro paragraph with **bold** and *italic*.

### Key Points

- First point
- Second point
- Third point

## Conclusion

Final thoughts here.
"""
        pdf = MarkdownPDF()
        pdf.render_markdown(md)
        buffer = pdf.output()
        assert len(buffer) > 500  # Non-trivial PDF


class TestMarkdownToPdf:
    @patch("pdf_tools.base_pdf_tool.upload_file")
    def test_markdown_to_pdf(self, mock_upload):
        mock_upload.return_value = "/output/test.pdf"
        result = markdown_to_pdf("# Hello\n\nWorld", file_name="test")
        assert result == "/output/test.pdf"
        mock_upload.assert_called_once()
        # Verify buffer is PDF
        args = mock_upload.call_args
        buffer = args[0][0]
        buffer.seek(0)
        assert buffer.read(4) == b"%PDF"

    @patch("pdf_tools.base_pdf_tool.upload_file")
    def test_markdown_to_pdf_no_name(self, mock_upload):
        mock_upload.return_value = "/output/random.pdf"
        result = markdown_to_pdf("Content here")
        assert result == "/output/random.pdf"


class TestDocxToPdf:
    @patch("pdf_tools.base_pdf_tool.upload_file")
    def test_docx_to_pdf(self, mock_upload, tmp_path):
        # Create sample docx
        path = str(tmp_path / "test.docx")
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Content paragraph.")
        doc.save(path)

        mock_upload.return_value = "/output/test.pdf"
        result = docx_to_pdf(path, file_name="test")
        assert result == "/output/test.pdf"
        # Verify PDF content
        args = mock_upload.call_args
        buffer = args[0][0]
        buffer.seek(0)
        assert buffer.read(4) == b"%PDF"

    @patch("pdf_tools.base_pdf_tool.upload_file")
    def test_docx_to_pdf_auto_name(self, mock_upload, tmp_path):
        path = str(tmp_path / "myreport.docx")
        doc = Document()
        doc.add_paragraph("Hello")
        doc.save(path)

        mock_upload.return_value = "/output/myreport.pdf"
        result = docx_to_pdf(path)
        # file_name should be "myreport" (from stem)
        args = mock_upload.call_args
        assert args[1].get("file_name") == "myreport" or args[0][2] == "myreport"

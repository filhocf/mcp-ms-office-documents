"""Tests for advanced DOCX features."""

import pytest
from docx import Document
from docx.shared import Cm
from pathlib import Path

from docx_tools.advanced import (
    add_image_to_docx, add_header_footer, set_page_margins,
    add_bullet_list, add_numbered_list, merge_table_cells,
)


@pytest.fixture
def sample_docx(tmp_path):
    path = str(tmp_path / "test.docx")
    doc = Document()
    doc.add_heading("Test Doc", level=1)
    doc.add_paragraph("Content here.")
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"R{i}C{j}"
    doc.save(path)
    return path


@pytest.fixture
def sample_image(tmp_path):
    """Create a minimal valid PNG image."""
    import struct
    import zlib

    path = tmp_path / "test.png"
    # Minimal 1x1 red PNG
    def create_png():
        signature = b'\x89PNG\r\n\x1a\n'
        # IHDR
        ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xffffffff
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
        # IDAT
        raw = b'\x00\xff\x00\x00'  # filter byte + RGB
        compressed = zlib.compress(raw)
        idat_crc = zlib.crc32(b'IDAT' + compressed) & 0xffffffff
        idat = struct.pack('>I', len(compressed)) + b'IDAT' + compressed + struct.pack('>I', idat_crc)
        # IEND
        iend_crc = zlib.crc32(b'IEND') & 0xffffffff
        iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
        return signature + ihdr + idat + iend

    path.write_bytes(create_png())
    return str(path)


class TestAddImage:
    def test_add_image(self, sample_docx, sample_image, tmp_path):
        out = str(tmp_path / "out.docx")
        result = add_image_to_docx(sample_docx, sample_image, width_inches=2.0, output_path=out)
        assert "Image added" in result
        assert Path(out).exists()


class TestHeaderFooter:
    def test_add_header(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        result = add_header_footer(sample_docx, header_text="My Header", output_path=out)
        assert "header" in result
        doc = Document(out)
        assert doc.sections[0].header.paragraphs[0].text == "My Header"

    def test_add_footer(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        result = add_header_footer(sample_docx, footer_text="Page Footer", output_path=out)
        assert "footer" in result
        doc = Document(out)
        assert doc.sections[0].footer.paragraphs[0].text == "Page Footer"

    def test_add_both(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        result = add_header_footer(sample_docx, header_text="H", footer_text="F", output_path=out)
        assert "header" in result and "footer" in result


class TestPageMargins:
    def test_set_margins(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        set_page_margins(sample_docx, top_cm=3.0, bottom_cm=3.0, left_cm=2.0, right_cm=2.0, output_path=out)
        doc = Document(out)
        section = doc.sections[0]
        assert abs(section.top_margin - Cm(3.0)) < Cm(0.01)
        assert abs(section.left_margin - Cm(2.0)) < Cm(0.01)


class TestLists:
    def test_bullet_list(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        result = add_bullet_list(sample_docx, ["Item A", "Item B", "Item C"], out)
        assert "3 items" in result
        doc = Document(out)
        # Check last 3 paragraphs are list items
        list_paras = [p for p in doc.paragraphs if "List" in (p.style.name or "")]
        assert len(list_paras) == 3

    def test_numbered_list(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        result = add_numbered_list(sample_docx, ["First", "Second"], out)
        assert "2 items" in result


class TestMergeTableCells:
    def test_merge_cells(self, sample_docx, tmp_path):
        out = str(tmp_path / "out.docx")
        result = merge_table_cells(sample_docx, 0, 0, 0, 0, 1, out)
        assert "Merged" in result
        doc = Document(out)
        # After merge, cell(0,0) and cell(0,1) should be same
        table = doc.tables[0]
        assert table.cell(0, 0)._tc == table.cell(0, 1)._tc

    def test_merge_invalid_table(self, sample_docx):
        with pytest.raises(ValueError):
            merge_table_cells(sample_docx, 99, 0, 0, 0, 1)

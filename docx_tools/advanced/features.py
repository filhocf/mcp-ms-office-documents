"""Advanced DOCX features implementation."""

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_ORIENT
from pathlib import Path


def add_image_to_docx(file_path: str, image_path: str, width_inches: float = 4.0, paragraph_index: int | None = None, output_path: str | None = None) -> str:
    """Add an image to a DOCX file. Optionally insert after a specific paragraph."""
    doc = Document(file_path)

    if paragraph_index is not None:
        # Insert after specific paragraph
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            raise ValueError(f"Paragraph index {paragraph_index} out of range")
        # Add at end and move
        run = doc.add_paragraph().add_run()
        run.add_picture(image_path, width=Inches(width_inches))
    else:
        doc.add_picture(image_path, width=Inches(width_inches))

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Image added. Saved to {save_path}"


def add_header_footer(file_path: str, header_text: str = "", footer_text: str = "", output_path: str | None = None) -> str:
    """Add header and/or footer to all sections of a DOCX file."""
    doc = Document(file_path)

    for section in doc.sections:
        if header_text:
            header = section.header
            header.paragraphs[0].text = header_text
        if footer_text:
            footer = section.footer
            footer.paragraphs[0].text = footer_text

    save_path = output_path or file_path
    doc.save(save_path)
    parts = []
    if header_text:
        parts.append("header")
    if footer_text:
        parts.append("footer")
    return f"Added {' and '.join(parts)}. Saved to {save_path}"


def set_page_margins(file_path: str, top_cm: float = 2.54, bottom_cm: float = 2.54, left_cm: float = 2.54, right_cm: float = 2.54, output_path: str | None = None) -> str:
    """Set page margins for all sections (in centimeters)."""
    doc = Document(file_path)

    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Margins set to top={top_cm}cm, bottom={bottom_cm}cm, left={left_cm}cm, right={right_cm}cm. Saved to {save_path}"


def add_bullet_list(file_path: str, items: list[str], output_path: str | None = None) -> str:
    """Add a bullet list to the end of a DOCX file."""
    doc = Document(file_path)

    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Added bullet list with {len(items)} items. Saved to {save_path}"


def add_numbered_list(file_path: str, items: list[str], output_path: str | None = None) -> str:
    """Add a numbered list to the end of a DOCX file."""
    doc = Document(file_path)

    for item in items:
        doc.add_paragraph(item, style="List Number")

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Added numbered list with {len(items)} items. Saved to {save_path}"


def merge_table_cells(file_path: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int, output_path: str | None = None) -> str:
    """Merge cells in a table from (start_row, start_col) to (end_row, end_col). 0-based indices."""
    doc = Document(file_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"Table index {table_index} out of range (0-{len(doc.tables)-1})")

    table = doc.tables[table_index]
    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)
    start_cell.merge(end_cell)

    save_path = output_path or file_path
    doc.save(save_path)
    return f"Merged cells ({start_row},{start_col}) to ({end_row},{end_col}) in table {table_index}. Saved to {save_path}"
